from docx import Document
from src.schema import ExecutionResult
from pathlib import Path

cwd = Path.cwd()
output_folder = cwd/"generated"

def generate_docx(
        result: ExecutionResult,
        filename: str
):
    doc = Document()

    for section in result.document:
        doc.add_heading(section.heading, level=1)

        for paragraph in section.paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph)
    
    filename = output_folder/filename
    doc.save(filename)
    return filename